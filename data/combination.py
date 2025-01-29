import subprocess
import pytesseract
from PIL import Image
from docx import Document
import os


def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""

    # Extract text from paragraphs and tables
    for para in doc.paragraphs:
        text += para.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"

    return text


# Step 2: Convert .docx to LaTeX using Pandoc
def convert_docx_to_latex(docx_file, latex_file):
    command = ['pandoc', docx_file, '-o', latex_file]
    subprocess.run(command, check=True)
    print(f"Converted {docx_file} to {latex_file}")


# Step 3: OCR for Image-based Formulas
def extract_formulas_from_image(image_file):
    # Open the image file
    img = Image.open(image_file)

    # Perform OCR on the image
    text = pytesseract.image_to_string(img, config='--psm 6')
    return text


# Step 4: Save output to files
def save_to_file(filename, content):
    with open(filename, 'w', encoding='utf-8') as file:
        file.write(content)


# Main function to run all steps
def process_docx_with_formulas(docx_file, image_file=None):
    # Step 1: Extract text from .docx file
    print("Extracting text from .docx file...")
    extracted_text = extract_text_from_docx(docx_file)
    save_to_file("extracted_text.txt", extracted_text)
    print("Text extracted and saved to 'extracted_text.txt'.")

    # Step 2: Convert .docx to LaTeX
    latex_file = "extracted_latex.tex"
    print("Converting .docx to LaTeX...")
    convert_docx_to_latex(docx_file, latex_file)
    print(f"LaTeX content saved to '{latex_file}'.")

    # Step 3: Optional - Extract formulas from image using OCR (if provided)
    if image_file and os.path.exists(image_file):
        print("Performing OCR on image to extract formulas...")
        formula_text = extract_formulas_from_image(image_file)
        save_to_file("extracted_formulas.txt", formula_text)
        print("Formulas extracted from image and saved to 'extracted_formulas.txt'.")
    else:
        print("No image file provided for OCR.")

    print("Processing complete!")


# Provide the .docx file and optional image file for OCR
docx_file = "Fizika.docx"  # Change this to your docx file
image_file = "image.png"  # Change this if you have an image with formulas

# Run the processing function
process_docx_with_formulas(docx_file, image_file)
